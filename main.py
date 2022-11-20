"""
# My first app
Here's our first attempt at using data to create a table:
"""

from fpdf import FPDF
from docx import Document
import numpy as np
from docx.shared import RGBColor
import fpdf


def write_docx(json_data, output_filename, highlight_correct=False):

    document = Document()
    ans_idxes = ['a', 'b', 'c', 'd']

    for key_idx in json_data.keys():

        p = document.add_paragraph(json_data[key_idx]['Question'])

        random_correct = [np.random.choice(json_data[key_idx]['Correct_answers'])]

        answers = np.random.choice(np.append(random_correct, json_data[key_idx]['False_answers']),
                                   np.clip(len(json_data[key_idx]['False_answers']) + 1, 2, 4), replace=False)

        for i, answer in enumerate(answers):
            if highlight_correct and (answer == random_correct[0]):
                run = p.add_run(f'\n\t{ans_idxes[i]}) {answer}')
                run.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
            else:
                run = p.add_run(f'\n\t{ans_idxes[i]}) {answer}')
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    document.add_page_break()

    docx_fname = output_filename + '.docx'

    document.save(docx_fname)

    


def write_pdf(json_data, output_filename, highlight_correct=False):

    pdf = FPDF()

    ans_idxes = ['a', 'b', 'c', 'd']

    pdf.add_page()

    pdf.set_font("Arial", size=15)

    help(pdf.set_font)
    help(fpdf.fpdf.FPDF)

    for key_idx in json_data.keys():
        pdf.set_font("Arial", size=15)
        pdf.cell(200, 10, txt=json_data[key_idx]['Question'],
                 ln=2, align='L')

        random_correct = [np.random.choice(json_data[key_idx]['Correct_answers'])]

        answers = np.random.choice(np.append(random_correct, json_data[key_idx]['False_answers']),
                                   np.clip(len(json_data[key_idx]['False_answers']) + 1, 2, 4), replace=False)

        pdf.set_font("Arial", size=10)
        for i, answer in enumerate(answers):
            if answer == random_correct[0] and highlight_correct:
                pdf.set_text_color(0, 255, 0)
                pdf.cell(40, 7, txt=f'\n{ans_idxes[i]}) {answer}',
                         ln=2, align='L')
                pdf.set_text_color(0, 0, 0)
            else:
                pdf.cell(40, 7, txt=f'\n{ans_idxes[i]}) {answer}',
                         ln=2, align='L')
    pdf_fname = output_filename + '.pdf'

    pdf.output(pdf_fname)

import streamlit as st
from gpt import *
# from export_functions import write_docx, write_pdf

### to delete!
sample_dict = {'1': {'Question': "Testowe pytanie?",
                         'Correct_answers': ["Poprawna :)"],
                         'False_answers': ["Nie poprawna 1", "Nie poprawna 2", "Nie poprawna 3"]},
                   '2': {'Question': "Testowe pytanie 2?",
                         'Correct_answers': ["Poprawna :)"],
                         'False_answers': ["Nie poprawna 1", "Nie poprawna 2", "Nie poprawna 3"]},
                   '3': {'Question': "Testowe pytanie 3?",
                         'Correct_answers': ["Tak"],
                         'False_answers': ["Nie"]}

                   }
### to delete!!!

def download_summarized_article(text: str) -> str:
    article = ""
    with st.spinner('Skracanie tekstu...'):
        article = summarize_article(text=text)
    return article

# streamlit_app.py


def check_password():
    """Returns `True` if the user had a correct password."""
    
    def password_entered():
        """Checks whether a password entered by the user is correct."""
        if (
            st.session_state["username"] in st.secrets["passwords"]
            and st.session_state["password"]
            == st.secrets["passwords"][st.session_state["username"]]
        ):
            st.session_state["password_correct"] = True
            del st.session_state["password"]  # don't store username + password
            del st.session_state["username"]
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        # First run, show inputs for username + password.
        st.text_input("Username", on_change=password_entered, key="username")
        st.text_input(
            "Password", type="password", on_change=password_entered, key="password"
        )
        return False
    elif not st.session_state["password_correct"]:
        # Password not correct, show input + error.
        st.text_input("Username", on_change=password_entered, key="username")
        st.text_input(
            "Password", type="password", on_change=password_entered, key="password"
        )
        st.error("😕 User not known or password incorrect")
  
        return False
    else:
        # Password correct.    
        return True
    
if check_password():

    st.write("# QuickQuiz")
    user_input = st.text_area("Tekst",placeholder="Wpisz tekst do wygenerowania quizu", label_visibility="hidden")

    if len(user_input) > 10:
        text = user_input
        # Wygeneruj pytania
        with st.spinner("Generuję pytania. Proszę czekać 🤔"):
            questions = generate_questions(text=text)
            st.write("Mam! 😋 Oto one: ")
            for i, question in enumerate(questions):
                st.write(f"**{i+1}. {question.strip()}**")

        # Wygeneruj odpowiedzi
        correct_answers: dict[int, list[str]] = dict()
        wrong_answers: dict[int, list[str]] = dict()
        with st.spinner("Generuję odpowiedzi 🏃🏼‍♂️🏃🏼‍♂️🏃🏼‍♂️"):
            for i, question in enumerate(questions):
                correct_answers[i] = generate_correct_answers(text=text, question=question)
                wrong_answers[i] = generate_wrong_answers(text=text, question=question)

        # Wypisz wyniki
        with st.spinner("Skończyłem 🥳🎉 oto wyniki:"):
            for i, question in enumerate(questions):
                st.write(f"## {i+1}. {question}")
                st.write(f"**Poprawne odpowiedzi:**")
                for j, ans in enumerate(correct_answers[i]):
                    st.write(f"\t{j+1}. {ans}")
                st.write(f"**Błędne odpowiedzi:**")
                for j, ans in enumerate(wrong_answers[i]):
                    st.write(f"\t{j+1}. {ans}")

    if st.button('Eksport pytań do pliku pdf.'):
        write_pdf(sample_dict, '2', highlight_correct=True)
        st.write('Wygenerowano plik pdf.')

    if st.button('Eksport pytań do pliku docx.'):
        write_docx(sample_dict, '2', highlight_correct=True)
        st.write('Wygenerowano plik docx.')


from fpdf import FPDF
from docx import Document
import numpy as np
from docx.shared import RGBColor
import fpdf


def write_docx(json_data, output_filename, highlight_correct=False):

    document = Document()
    ans_idxes = ['a', 'b', 'c', 'd']

    for key_idx in json_data.keys():

        p = document.add_paragraph(json_data[key_idx]['Question'])

        random_correct = [np.random.choice(json_data[key_idx]['Correct_answers'])]

        answers = np.random.choice(np.append(random_correct, json_data[key_idx]['False_answers']),
                                   np.clip(len(json_data[key_idx]['False_answers']) + 1, 2, 4), replace=False)

        for i, answer in enumerate(answers):
            if highlight_correct and (answer == random_correct[0]):
                run = p.add_run(f'\n\t{ans_idxes[i]}) {answer}')
                run.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
            else:
                run = p.add_run(f'\n\t{ans_idxes[i]}) {answer}')
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    document.add_page_break()

    docx_fname = output_filename + '.docx'

    document.save(docx_fname)

    


def write_pdf(json_data, output_filename, highlight_correct=False):

    pdf = FPDF()

    ans_idxes = ['a', 'b', 'c', 'd']

    pdf.add_page()

    pdf.set_font("Arial", size=15)

    help(pdf.set_font)
    help(fpdf.fpdf.FPDF)

    for key_idx in json_data.keys():
        pdf.set_font("Arial", size=15)
        pdf.cell(200, 10, txt=json_data[key_idx]['Question'],
                 ln=2, align='L')

        random_correct = [np.random.choice(json_data[key_idx]['Correct_answers'])]

        answers = np.random.choice(np.append(random_correct, json_data[key_idx]['False_answers']),
                                   np.clip(len(json_data[key_idx]['False_answers']) + 1, 2, 4), replace=False)

        pdf.set_font("Arial", size=10)
        for i, answer in enumerate(answers):
            if answer == random_correct[0] and highlight_correct:
                pdf.set_text_color(0, 255, 0)
                pdf.cell(40, 7, txt=f'\n{ans_idxes[i]}) {answer}',
                         ln=2, align='L')
                pdf.set_text_color(0, 0, 0)
            else:
                pdf.cell(40, 7, txt=f'\n{ans_idxes[i]}) {answer}',
                         ln=2, align='L')
    pdf_fname = output_filename + '.pdf'

    pdf.output(pdf_fname)