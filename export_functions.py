# pip install fpdf
# pip install docx
# pip install numpy

from fpdf import FPDF
from docx import Document
import numpy as np
from docx.shared import RGBColor
import fpdf


def write_docx(json_data, output_filename, highlight_correct=False):

    document = Document()
    ans_idxes = ['a', 'b', 'c', 'd']

    for key_idx in json_data.keys():

        p = document.add_paragraph(json_data[key_idx]['Question'])

        random_correct = json_data[key_idx]['Correct_answers']

        answers = np.random.choice(np.append(random_correct, json_data[key_idx]['False_answers'][:(4 - len(random_correct))]),
                                   np.clip(len(json_data[key_idx]['False_answers']) + len(random_correct), 0, 4), replace=False)

        for i, answer in enumerate(answers):
            if np.isin(answer, random_correct) and highlight_correct:
                run = p.add_run(f'\n\t{ans_idxes[i]}) {answer}')
                run.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
            else:
                run = p.add_run(f'\n\t{ans_idxes[i]}) {answer}')
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    document.add_page_break()

    docx_fname = output_filename + '.docx'

    document.save(docx_fname)


def write_pdf(json_data, output_filename, highlight_correct=False):

    pdf = FPDF()

    ans_idxes = ['a', 'b', 'c', 'd']

    pdf.add_page()

    pdf.set_font("Arial", size=15)

    for key_idx in json_data.keys():
        pdf.set_font("Arial", size=15)
        pdf.cell(200, 10, txt=pdf.normalize_text(json_data[key_idx]['Question'].encode('latin-1', 'ignore').decode('latin-1', 'ignore')),
                 ln=2, align='L')

        random_correct = json_data[key_idx]['Correct_answers']

        answers = np.random.choice(np.append(random_correct, json_data[key_idx]['False_answers'][:(4 - len(random_correct))]),
                                   np.clip(len(json_data[key_idx]['False_answers']) + len(random_correct), 0, 4), replace=False)

        pdf.set_font("Arial", size=10)
        for i, answer in enumerate(answers):
            if np.isin(answer, random_correct) and highlight_correct:
                pdf.set_text_color(0, 255, 0)
                pdf.cell(40, 7, txt=pdf.normalize_text(f"\n{ans_idxes[i]}) {answer.encode('latin-1', 'ignore').decode('latin-1', 'ignore')}"),
                         ln=2, align='L')
                pdf.set_text_color(0, 0, 0)
            else:
                pdf.cell(40, 7, txt=pdf.normalize_text(f"\n{ans_idxes[i]}) {answer.encode('latin-1', 'ignore').decode('latin-1', 'ignore')}"),
                         ln=2, align='L')
    pdf_fname = output_filename + '.pdf'

    pdf.output(pdf_fname)


if __name__ == '__main__':

    """
    sample_dict = {'1': {'Question': "Testowe pytanie?",
                         'Correct_answers': ["Poprawna :)"],
                         'False_answers': ["Nie poprawna 1", "Nie poprawna 2", "Nie poprawna 3"]},
                   '2': {'Question': "Testowe pytanie 2?",
                         'Correct_answers': ["Poprawna :)"],
                         'False_answers': ["Nie poprawna 1", "Nie poprawna 2", "Nie poprawna 3"]},
                   '3': {'Question': "Testowe pytanie 3?",
                         'Correct_answers': ["Tak"],
                         'False_answers': ["Nie"]}

                   }
    """

    sample_dict = {1: {'Question': 'W jakim czasie zako??czy??a si?? II wojna ??wiatowa? ',
                       'Correct_answers': ['1949, 1955, 1979.'],
                       'False_answers': ['2 wrze??nia 1945 roku ??? kapitulacja Japonii.', '9 maja 1945 roku ??? wej??cie w ??ycie aktu bezwarunkowej kapitulacji III Rzeszy.', '8 maja 1945 roku ??? podpisanie aktu bezwarunkowej kapitulacji III Rzeszy.']},
                   2: {'Question': 'Czy by??a to jedna wielka wojna?', 'Correct_answers': ['Wojna by??a wielk??, ale nie jedn?? wielk?? wojn??.', 'Wojna by??a wielk??, ale nie jedn?? wielk?? wojn?? ojczy??nian??.', 'Wojna by??a wielk??, ale nie jedn?? wielk?? wojn?? ??wiatow??.'],
                       'False_answers': ['      ', 'by??a to jedna wielka wojna, ale liczba stron mia??a inn?? skal??.', 'by??a to jedna wielka wojna, ale r????nice w liczebno??ciach ??o??nierzy i ofiar by??y du??e;', 'by??a to jedna wielka wojna, cho?? podzia?? na strony mia?? inn?? skal??;']},
                   3: {'Question': 'Jak?? liczb?? ludzi zgin????o w II wojnie ??wiatowej? ', 'Correct_answers': [' 1,7 mln ludzi zgin????o w II wojnie ??wiatowej.', '2,0 mln ludzi zgin????o w II wojnie ??wiatowej.', '3,0 mln ludzi zgin????o w II wojnie ??wiatowej.'],
                       'False_answers': ['Odpowied?? 2:', '  Wszystkie szacunki na ten temat s?? nieprawdziwe.', 'Odpowied?? 3:', 'Wed??ug szacunk??w zgin????o w niej od 50 do 78 milion??w ludzi.', 'Odpowied?? 1:']},
                   4: {'Question': 'Czy wojna by??a trwa??a dwukrotnie w Europie? ', 'Correct_answers': ['Tak, wojna trwa??a dwukrotnie w Europie.', 'Tak, wojna by??a trwa??a dwukrotnie w Europie.'],
                       'False_answers': ['     ', 'Wed??ug r????nych szacunk??w zgin????o w niej od 50[2] do 78[3] milion??w ludzi.']},
                   5: {'Question': 'Jak?? rol?? wojennej w II wojnie ??wiatowej odgrywa??a Europa? ',
                       'Correct_answers': [' Europa by??a g????wnym obszarem dzia??a?? wojennych w II wojnie ??wiatowej. Wojna obj????a prawie ca???? Europ??, wschodni?? i po??udniowo-wschodni?? Azj??, p????nocn?? Afryk??, cz?????? Bliskiego Wschodu, wyspy Oceanii i wszystkie oceany. Do walki przyst??pi??o 1,7 mld ludzi.'],
                       'False_answers': ['Europa by??a g????wnym cz????ci?? obszaru dzia??a?? wojennych wojny.', '  1. Europa by??a jednym z g????wnych obszar??w dzia??a?? wojennych wojny.', 'Europa by??a g????wnym stron?? konfliktu.']},
                   6: {'Question': 'Czy zgin????o w niej wszystkich, czy tylko cz?????? ludzi? ',
                       'Correct_answers': ['Wszystkich zgin????o.', 'Cz?????? ludzi zgin????a, ale wszystkich nie zabi??a. ', ' Cz?????? ludzi zgin????a, ale wszystkich nie zabi??a. '],
                       'False_answers': [' 1. Wed??ug r????nych szacunk??w zgin????o w niej od 50[2] do 78[3] milion??w ludzi.', 'Wed??ug r????nych szacunk??w zgin????o w niej od 50[2] do 78[3] milion??w ludzi.']}}


    write_docx(sample_dict, '2', highlight_correct=True)
    write_pdf(sample_dict, '2', highlight_correct=True)
